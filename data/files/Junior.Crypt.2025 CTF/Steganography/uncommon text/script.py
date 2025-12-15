from docx import Document

# Визуально одинаковые русские и английские буквы
visually_similar_letters = {
    'A': ('А', 'A'), 'a': ('а', 'a'),
    'B': ('В', 'B'),
    'C': ('С', 'C'), 'c': ('с', 'c'),
    'E': ('Е', 'E'), 'e': ('е', 'e'),
    'H': ('Н', 'H'),
    'K': ('К', 'K'),
    'M': ('М', 'M'),
    'O': ('О', 'O'), 'o': ('о', 'o'),
    'P': ('Р', 'P'), 'p': ('р', 'p'),
    'T': ('Т', 'T'),
    'X': ('Х', 'X'), 'x': ('х', 'x'),
    'Y': ('У', 'Y'), 'y': ('у', 'y'),
}

# Для встраивания
letter_to_pair = {ru: (ru, en) for _, (ru, en) in visually_similar_letters.items()}

# Для извлечения
all_target_chars = {}
for _, (ru, en) in visually_similar_letters.items():
    all_target_chars[ru] = '0'
    all_target_chars[en] = '1'

# Преобразует текст флага в двоичный код
def text_to_binary(flag):
    return ''.join(f'{ord(c):08b}' for c in flag)

# Читает весь текст из docx
def read_text_from_docx(path):
    doc = Document(path)
    return '\n'.join([para.text for para in doc.paragraphs])

# Записывает текст в docx
def write_text_to_docx(text, path):
    doc = Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    doc.save(path)

# Встраивает флаг в текст
def embed_flag_in_text(text, flag):
    binary = text_to_binary(flag)
    binary_index = 0
    result = ''

    for char in text:
        if char in letter_to_pair and binary_index < len(binary):
            ru, en = letter_to_pair[char]
            bit = binary[binary_index]
            result += ru if bit == '0' else en
            binary_index += 1
        else:
            result += char

    if binary_index < len(binary):
        print(f"[!] Недостаточно подходящих букв: нужно {len(binary)}, найдено {binary_index}")
    else:
        print(f"[+] Флаг успешно встроен ({binary_index} бит).")

    return result

# Извлекает бинарный флаг из docx
def extract_flag_from_docx(path, flag_length_bits):
    doc = Document(path)
    binary_string = ''

    for para in doc.paragraphs:
        for char in para.text:
            if char in all_target_chars:
                binary_string += all_target_chars[char]

    if len(binary_string) < flag_length_bits:
        print(f"[!] Недостаточно данных для флага ({len(binary_string)} из {flag_length_bits})")
        return None

    flag_bits = binary_string[:flag_length_bits]
    flag_bytes = [flag_bits[i:i+8] for i in range(0, len(flag_bits), 8)]

    try:
        return ''.join(chr(int(b, 2)) for b in flag_bytes)
    except:
        print("[!] Ошибка при декодировании.")
        return None

# === Основной запуск ===
if __name__ == "__main__":
    input_path = input("Введите путь к исходному .docx файлу: ").strip()
    flag = input("Введите флаг, который нужно скрыть: ").strip()
    output_path = input("Введите путь для сохранения нового .docx файла: ").strip()
    # Читаем оригинальный текст
    text = read_text_from_docx(input_path)

    # Встраиваем флаг
    embedded = embed_flag_in_text(text, flag)

    # Сохраняем в новый файл
    write_text_to_docx(embedded, output_path)
    print(f"[+] Готово. Файл с флагом сохранён: {output_path}")


